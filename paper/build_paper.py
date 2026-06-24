"""Build the IEEE-formatted paper (English + Chinese) as .docx files, reusing
the supplied IEEE Transactions template for all paragraph styles and the
single-column-masthead / two-column-body section layout. Equations are real
OMML (Word native math); body text is justified.

    python3 -m paper.build_paper
"""
import os
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Inches, Pt, Emu
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn

from paper.omml import omath
from paper.equations import build as build_eqs
from paper.inline import compile_inline, split_text
from paper import content_en, content_zh

HERE = os.path.dirname(__file__)
ROOT = os.path.dirname(HERE)
TEMPLATE = ("/home/eugene/.claude2/uploads/16919b89-26fc-4278-8786-f30605a38ba5/"
            "9f8f78aa-Transactionstemplateandinstructionsonhowtocreateyourarticle"
            "formatted_1.docx")
COL_W = 3.3   # inches, single body column width for figures / eq tab


# --------------------------------------------------------------------------- #
def set_cols(section, num, space=360):
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(space))
    if num > 1:
        cols.set(qn("w:equalWidth"), "1")


def set_cjk_default(doc, font="PMingLiU"):
    """Set the document-default East-Asian font so Chinese renders correctly."""
    styles_el = doc.styles.element
    dd = styles_el.find(qn("w:docDefaults"))
    if dd is None:
        dd = OxmlElement("w:docDefaults")
        styles_el.insert(0, dd)
    rprd = dd.find(qn("w:rPrDefault"))
    if rprd is None:
        rprd = OxmlElement("w:rPrDefault")
        dd.append(rprd)
    rpr = rprd.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        rprd.append(rpr)
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), font)


class Paper:
    def __init__(self, template, cjk_font=None):
        self.doc = Document(template)
        self.eqs = build_eqs()
        self._clear_body()
        if cjk_font:
            set_cjk_default(self.doc, cjk_font)
        self._masthead_done = False
        # IEEE body: first paragraph after a heading is flush-left, every
        # subsequent paragraph gets a first-line indent (matches the template).
        self._after_heading = True

    # -- low level ------------------------------------------------------- #
    def _clear_body(self):
        body = self.doc.element.body
        for child in list(body):
            if child.tag == qn("w:sectPr"):
                continue
            body.remove(child)

    def _style(self, name, fallback="Normal"):
        names = {s.name for s in self.doc.styles}
        return name if name in names else fallback

    def _p(self, style="Normal", align=None):
        p = self.doc.add_paragraph(style=self._style(style))
        if align is not None:
            p.alignment = align
        return p

    # -- masthead -------------------------------------------------------- #
    def title(self, text):
        p = self._p("Title", WD_ALIGN_PARAGRAPH.CENTER)
        p.add_run(text)

    def authors(self, text):
        p = self._p("Normal", WD_ALIGN_PARAGRAPH.CENTER)
        p.add_run(text)

    def abstract(self, lead, text):
        p = self._p(self._style("Abstract", "Normal"), WD_ALIGN_PARAGRAPH.JUSTIFY)
        r = p.add_run(lead)
        r.bold = True
        p.add_run(text)

    def index_terms(self, lead, text):
        p = self._p(self._style("IndexTerms", "Normal"),
                    WD_ALIGN_PARAGRAPH.JUSTIFY)
        r = p.add_run(lead)
        r.bold = True
        r.italic = True
        p.add_run(text)

    def start_body(self):
        """Close the 1-column masthead section, open the 2-column body."""
        self.doc.add_section(WD_SECTION.CONTINUOUS)
        set_cols(self.doc.sections[0], 1)
        set_cols(self.doc.sections[1], 2)

    # -- body ------------------------------------------------------------ #
    def h1(self, text):
        self._p("Heading 1").add_run(text)
        self._after_heading = True

    def h2(self, text):
        self._p("Heading 2").add_run(text)
        self._after_heading = True

    def para(self, text):
        p = self._p("Normal", WD_ALIGN_PARAGRAPH.JUSTIFY)
        # Every body paragraph gets a first-line indent (incl. the first one of
        # each section).
        self._after_heading = False
        p.paragraph_format.first_line_indent = Emu(128270)  # ~0.14 in
        # Render $...$ spans as genuine inline Word math; plain text otherwise.
        for is_math, chunk in split_text(text):
            if is_math:
                p._p.append(parse_xml(omath(compile_inline(chunk))))
            else:
                p.add_run(chunk)
        return p

    def equation(self, number):
        inner = self.eqs[number]
        p = self._p(self._style("Equation", "Normal"))
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Emu(0)
        # centre tab carries the formula, right tab carries the number (IEEE)
        pf = p.paragraph_format
        pf.tab_stops.add_tab_stop(Inches(COL_W / 2), WD_TAB_ALIGNMENT.CENTER)
        pf.tab_stops.add_tab_stop(Inches(COL_W), WD_TAB_ALIGNMENT.RIGHT)
        p.add_run("\t")
        p._p.append(parse_xml(omath(inner)))
        p.add_run(f"\t({number})")

    def figure(self, path, caption, width=COL_W):
        p = self._p("Normal", WD_ALIGN_PARAGRAPH.CENTER)
        p.add_run().add_picture(path, width=Inches(width))
        c = self._p(self._style("Figure Caption", "Normal"),
                    WD_ALIGN_PARAGRAPH.CENTER)
        c.add_run(caption)

    def wide_figure(self, path, caption, width=6.9):
        """Full-width (both-column-spanning) figure via a 1-column band."""
        self.doc.add_section(WD_SECTION.CONTINUOUS)
        set_cols(self.doc.sections[-1], 1)          # band is single column
        self.figure(path, caption, width=width)
        self.doc.add_section(WD_SECTION.CONTINUOUS)
        set_cols(self.doc.sections[-1], 2)          # resume two columns

    def table(self, label, caption, header, rows):
        self._p(self._style("Table Title", "Normal"),
                WD_ALIGN_PARAGRAPH.CENTER).add_run(label)
        self._p(self._style("Table Title", "Normal"),
                WD_ALIGN_PARAGRAPH.CENTER).add_run(caption)
        tbl = self.doc.add_table(rows=1 + len(rows), cols=len(header))
        try:
            tbl.style = "Table Grid"
        except Exception:
            pass
        tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for j, htxt in enumerate(header):
            cell = tbl.rows[0].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(htxt)
            run.bold = True
            run.font.size = Pt(8)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, row in enumerate(rows, start=1):
            for j, val in enumerate(row):
                cell = tbl.rows[i].cells[j]
                cell.text = ""
                run = cell.paragraphs[0].add_run(str(val))
                run.font.size = Pt(8)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._p("Normal")  # spacer

    def algorithm_box(self, title, lines):
        """A bordered pseudocode box (single full-width cell, monospace)."""
        self._p(self._style("Table Title", "Normal"),
                WD_ALIGN_PARAGRAPH.LEFT).add_run(title)
        tbl = self.doc.add_table(rows=1, cols=1)
        try:
            tbl.style = "Table Grid"
        except Exception:
            pass
        cell = tbl.rows[0].cells[0]
        cell.text = ""
        for i, ln in enumerate(lines):
            par = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            par.paragraph_format.left_indent = Inches(0.05)
            run = par.add_run(ln)
            run.font.name = "Consolas"
            run.font.size = Pt(8)
        self._p("Normal")

    def references(self, items):
        for it in items:
            p = self._p(self._style("References", "Normal"),
                        WD_ALIGN_PARAGRAPH.JUSTIFY)
            p.add_run(it)

    def save(self, path):
        self.doc.save(path)


def main():
    for tag, mod, cjk, out in [
        ("EN", content_en, None, os.path.join(ROOT, "paper_IEEE_EN.docx")),
        ("ZH", content_zh, "PMingLiU", os.path.join(ROOT, "paper_IEEE_ZH.docx")),
    ]:
        p = Paper(TEMPLATE, cjk_font=cjk)
        mod.build(p)
        p.save(out)
        print(f"saved {tag} -> {out}")


if __name__ == "__main__":
    main()
